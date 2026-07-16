"""File processing service for handling various document formats"""

import os
import io
import re
import chardet
from typing import Optional, Tuple, List
from pathlib import Path
import logging
from enum import Enum

# File type specific imports
import PyPDF2
from docx import Document as DocxDocument
import markdown
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)


class FileType(str, Enum):
    """Supported file types"""
    TXT = "txt"
    MD = "md"
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    HTM = "htm"
    CSV = "csv"
    JSON = "json"
    XML = "xml"
    PY = "py"
    JS = "js"
    UNKNOWN = "unknown"


class FileProcessor:
    """
    Service for processing and extracting text from various file formats.
    
    Supported formats:
    - .txt: Plain text
    - .md: Markdown
    - .pdf: PDF documents
    - .docx: Word documents
    - .html/.htm: HTML pages
    - .csv: CSV files
    - .json: JSON files
    - .xml: XML files
    - .py: Python source code
    - .js: JavaScript source code
    """
    
    def __init__(self):
        """Initialize the file processor"""
        self.max_file_size = 10 * 1024 * 1024  # 10MB
        logger.info("File Processor initialized")
    
    def detect_file_type(self, filename: str) -> FileType:
        """
        Detect file type from extension.
        
        Args:
            filename: Name of the file
            
        Returns:
            FileType enum
        """
        extension = Path(filename).suffix.lower()
        
        file_type_map = {
            '.txt': FileType.TXT,
            '.md': FileType.MD,
            '.markdown': FileType.MD,
            '.pdf': FileType.PDF,
            '.docx': FileType.DOCX,
            '.html': FileType.HTML,
            '.htm': FileType.HTM,
            '.csv': FileType.CSV,
            '.json': FileType.JSON,
            '.xml': FileType.XML,
            '.py': FileType.PY,
            '.js': FileType.JS,
        }
        
        return file_type_map.get(extension, FileType.UNKNOWN)
    
    async def process_file(
        self,
        file_content: bytes,
        filename: str,
    ) -> Tuple[str, str, dict]:
        """
        Process a file and extract text content.
        
        Args:
            file_content: Raw file content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (extracted_text, file_type, metadata)
        """
        # Check file size
        if len(file_content) > self.max_file_size:
            raise ValueError(f"File too large: {len(file_content)} bytes (max: {self.max_file_size})")
        
        # Detect file type
        file_type = self.detect_file_type(filename)
        
        if file_type == FileType.UNKNOWN:
            # Try to detect from content
            file_type = self._detect_from_content(file_content)
        
        logger.info(f"Processing file: {filename} (type: {file_type})")
        
        # Process based on file type
        try:
            if file_type == FileType.TXT:
                text, metadata = await self._process_text(file_content)
            elif file_type == FileType.MD:
                text, metadata = await self._process_markdown(file_content)
            elif file_type == FileType.PDF:
                text, metadata = await self._process_pdf(file_content)
            elif file_type == FileType.DOCX:
                text, metadata = await self._process_docx(file_content)
            elif file_type in [FileType.HTML, FileType.HTM]:
                text, metadata = await self._process_html(file_content)
            elif file_type == FileType.CSV:
                text, metadata = await self._process_csv(file_content)
            elif file_type == FileType.JSON:
                text, metadata = await self._process_json(file_content)
            elif file_type == FileType.XML:
                text, metadata = await self._process_xml(file_content)
            elif file_type in [FileType.PY, FileType.JS]:
                text, metadata = await self._process_code(file_content, file_type)
            else:
                # Fallback: try to decode as text
                text, metadata = await self._process_text(file_content)
            
            # Clean and normalize text
            text = self._clean_text(text)
            
            # Add file metadata
            metadata.update({
                "filename": filename,
                "file_type": file_type.value,
                "file_size": len(file_content),
            })
            
            logger.info(f"✅ Processed file: {filename} ({len(text)} characters)")
            return text, file_type.value, metadata
            
        except Exception as e:
            logger.error(f"❌ Failed to process file {filename}: {e}")
            raise ValueError(f"Failed to process file: {str(e)}")
    
    def _detect_from_content(self, content: bytes) -> FileType:
        """Detect file type from content"""
        # Try to detect if it's text
        try:
            # Try UTF-8
            text = content.decode('utf-8')
            # Check for common patterns
            if '<html' in text.lower() or '<!doctype html' in text.lower():
                return FileType.HTML
            if '<?xml' in text[:100]:
                return FileType.XML
            if text.strip().startswith('{') or text.strip().startswith('['):
                try:
                    import json
                    json.loads(text)
                    return FileType.JSON
                except:
                    pass
        except:
            pass
        
        return FileType.TXT
    
    async def _process_text(self, content: bytes) -> Tuple[str, dict]:
        """Process plain text file"""
        try:
            # Detect encoding
            encoding = self._detect_encoding(content)
            text = content.decode(encoding, errors='ignore')
            
            metadata = {
                "encoding": encoding,
                "line_count": len(text.splitlines()),
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Text processing error: {e}")
            # Fallback to UTF-8 with error replacement
            text = content.decode('utf-8', errors='replace')
            return text, {"encoding": "utf-8", "errors": "replaced"}
    
    async def _process_markdown(self, content: bytes) -> Tuple[str, dict]:
        """Process markdown file"""
        try:
            # Decode content
            encoding = self._detect_encoding(content)
            text = content.decode(encoding, errors='ignore')
            
            # Convert markdown to plain text
            html = markdown.markdown(text)
            soup = BeautifulSoup(html, 'html.parser')
            plain_text = soup.get_text(separator='\n')
            
            metadata = {
                "markdown": True,
                "encoding": encoding,
                "original_text": text[:500],  # Store preview
            }
            
            return plain_text, metadata
            
        except Exception as e:
            logger.error(f"Markdown processing error: {e}")
            # Fallback to plain text
            encoding = self._detect_encoding(content)
            text = content.decode(encoding, errors='ignore')
            return text, {"markdown": False, "encoding": encoding}
    
    async def _process_pdf(self, content: bytes) -> Tuple[str, dict]:
        """Process PDF file"""
        try:
            text = ""
            metadata = {
                "pages": 0,
                "author": "",
                "title": "",
            }
            
            # Create a BytesIO object
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Extract metadata
            if pdf_reader.metadata:
                metadata["author"] = pdf_reader.metadata.get('/Author', '')
                metadata["title"] = pdf_reader.metadata.get('/Title', '')
                metadata["creator"] = pdf_reader.metadata.get('/Creator', '')
            
            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {page_num} ---\n{page_text}"
            
            metadata["pages"] = len(pdf_reader.pages)
            metadata["word_count"] = len(text.split())
            
            if not text.strip():
                raise ValueError("PDF appears to be scanned or contains no text")
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            raise ValueError(f"Failed to process PDF: {str(e)}")
    
    async def _process_docx(self, content: bytes) -> Tuple[str, dict]:
        """Process DOCX file"""
        try:
            # Create a BytesIO object
            docx_file = io.BytesIO(content)
            doc = DocxDocument(docx_file)
            
            # Extract text from paragraphs
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text += f"\n{row_text}"
            
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "word_count": len(text.split()),
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            raise ValueError(f"Failed to process DOCX: {str(e)}")
    
    async def _process_html(self, content: bytes) -> Tuple[str, dict]:
        """Process HTML file"""
        try:
            # Decode content
            encoding = self._detect_encoding(content)
            text = content.decode(encoding, errors='ignore')
            
            # Parse HTML and extract text
            soup = BeautifulSoup(text, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text
            plain_text = soup.get_text(separator='\n')
            
            # Extract title
            title = soup.title.string if soup.title else ""
            
            # Clean up whitespace
            plain_text = '\n'.join(line.strip() for line in plain_text.splitlines() if line.strip())
            
            metadata = {
                "html": True,
                "title": title.strip() if title else "",
                "encoding": encoding,
            }
            
            return plain_text, metadata
            
        except Exception as e:
            logger.error(f"HTML processing error: {e}")
            # Fallback to plain text
            encoding = self._detect_encoding(content)
            text = content.decode(encoding, errors='ignore')
            return text, {"html": False, "encoding": encoding}
    
    async def _process_csv(self, content: bytes) -> Tuple[str, dict]:
        """Process CSV file"""
        try:
            # Decode content
            encoding = self._detect_encoding(content)
            text = content.decode(encoding, errors='ignore')
            
            # Format CSV as readable text
            lines = text.splitlines()
            if not lines:
                return text, {"csv": True, "rows": 0}
            
            # Try to detect delimiter
            first_line = lines[0]
            delimiter = self._detect_csv_delimiter(first_line)
            
            # Format as text
            formatted_lines = []
            for i, line in enumerate(lines[:50]):  # Limit to 50 rows
                parts = line.split(delimiter)
                if i == 0:
                    formatted_lines.append(f"Headers: {' | '.join(parts)}")
                    formatted_lines.append("-" * 40)
                else:
                    formatted_lines.append(f"Row {i}: {' | '.join(parts)}")
            
            if len(lines) > 50:
                formatted_lines.append(f"... and {len(lines) - 50} more rows")
            
            formatted_text = "\n".join(formatted_lines)
            
            metadata = {
                "csv": True,
                "rows": len(lines),
                "delimiter": delimiter,
                "encoding": encoding,
            }
            
            return formatted_text, metadata
            
        except Exception as e:
            logger.error(f"CSV processing error: {e}")
            # Fallback to plain text
            encoding = self._detect_encoding(content)
            text = content.decode(encoding, errors='ignore')
            return text, {"csv": False, "encoding": encoding}
    
    async def _process_json(self, content: bytes) -> Tuple[str, dict]:
        """Process JSON file"""
        try:
            # First decode the content
            encoding = self._detect_encoding(content)
            text = content.decode(encoding, errors='ignore')
            
            # Pretty print JSON
            import json
            data = json.loads(text)
            formatted = json.dumps(data, indent=2)
            metadata = {
                "json": True,
                "pretty": True,
                "encoding": encoding,
            }
            return formatted, metadata
        except json.JSONDecodeError as e:
            logger.warning(f"JSON parsing failed: {e}, returning as text")
            # If JSON parsing fails, return as text
            encoding = self._detect_encoding(content)
            text = content.decode(encoding, errors='ignore')
            metadata = {
                "json": False,
                "encoding": encoding,
            }
            return text, metadata
        except Exception as e:
            logger.error(f"JSON processing error: {e}")
            # Fallback to text
            encoding = self._detect_encoding(content)
            text = content.decode(encoding, errors='ignore')
            metadata = {
                "json": False,
                "encoding": encoding,
            }
            return text, metadata
    
    async def _process_xml(self, content: bytes) -> Tuple[str, dict]:
        """Process XML file"""
        try:
            # First decode the content
            encoding = self._detect_encoding(content)
            text = content.decode(encoding, errors='ignore')
            
            # Format XML as readable text
            from xml.dom import minidom
            xml_doc = minidom.parseString(text)
            formatted = xml_doc.toprettyxml(indent="  ")
            metadata = {
                "xml": True,
                "pretty": True,
                "encoding": encoding,
            }
            return formatted, metadata
        except Exception as e:
            logger.warning(f"XML parsing failed: {e}, returning as text")
            # If XML parsing fails, return as text
            encoding = self._detect_encoding(content)
            text = content.decode(encoding, errors='ignore')
            metadata = {
                "xml": False,
                "encoding": encoding,
            }
            return text, metadata
    
    async def _process_code(self, content: bytes, file_type: FileType) -> Tuple[str, dict]:
        """Process code files"""
        text, metadata = await self._process_text(content)
        
        # Add code formatting indicators
        language = "python" if file_type == FileType.PY else "javascript"
        formatted = f"```{language}\n{text}\n```"
        
        metadata.update({
            "code": True,
            "language": language,
            "line_count": len(text.splitlines()),
        })
        
        return formatted, metadata
    
    def _detect_encoding(self, content: bytes) -> str:
        """Detect text encoding"""
        try:
            result = chardet.detect(content[:10000])  # Sample first 10KB
            encoding = result['encoding'] if result['encoding'] else 'utf-8'
            confidence = result['confidence'] if result['confidence'] else 0
            
            # If confidence is low, try UTF-8
            if confidence < 0.7:
                try:
                    content.decode('utf-8')
                    return 'utf-8'
                except:
                    pass
            
            return encoding
        except:
            return 'utf-8'
    
    def _detect_csv_delimiter(self, line: str) -> str:
        """Detect CSV delimiter"""
        common_delimiters = [',', ';', '\t', '|']
        delimiter_counts = {}
        
        for delim in common_delimiters:
            count = line.count(delim)
            if count > 0:
                delimiter_counts[delim] = count
        
        if delimiter_counts:
            return max(delimiter_counts, key=delimiter_counts.get)
        return ','
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Remove multiple blank lines
        text = re.sub(r' +', ' ', text)  # Remove multiple spaces
        
        # Remove control characters except newline and tab
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', text)
        
        return text.strip()


# Create a global instance
file_processor = FileProcessor()